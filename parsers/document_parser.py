"""
Document parser module.

Supports CSV, Excel (.xlsx/.xls), PDF, and Word (.docx) files.
Each parser returns a ``ParsedDocument`` dataclass that carries both the
raw text representation and, where applicable, a structured
``pandas.DataFrame`` for tabular data.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pandas as pd
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

# Maximum permitted file size (bytes). Files larger than this are rejected
# before any parsing library processes them, preventing memory exhaustion.
_MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB


@dataclass
class ParsedDocument:
    """Container for the output of a document parse operation."""

    filename: str
    file_type: str  # "csv" | "excel" | "pdf" | "word"
    text_content: str
    dataframe: Optional[pd.DataFrame] = field(default=None, repr=False)


class DocumentParser:
    """Parse uploaded files into a unified ``ParsedDocument`` representation.

    Supported formats
    -----------------
    * CSV  (``.csv``)
    * Excel (``.xlsx``, ``.xls``)
    * PDF  (``.pdf``)
    * Word (``.docx``)
    """

    # Map file extensions → internal type key
    _EXTENSION_MAP: dict[str, str] = {
        ".csv": "csv",
        ".xlsx": "excel",
        ".xls": "excel",
        ".pdf": "pdf",
        ".docx": "word",
    }

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse *file_bytes* coming from an uploaded file named *filename*.

        Parameters
        ----------
        file_bytes:
            Raw bytes of the uploaded file.
        filename:
            Original file name (used to detect the format).

        Returns
        -------
        ParsedDocument
            A populated ``ParsedDocument`` instance.

        Raises
        ------
        ValueError
            If the file extension is not supported or the file exceeds the
            maximum allowed size (50 MB).
        """
        if len(file_bytes) > _MAX_FILE_SIZE_BYTES:
            max_mb = _MAX_FILE_SIZE_BYTES // (1024 * 1024)
            raise ValueError(
                f"File '{filename}' is too large "
                f"({len(file_bytes) / (1024 * 1024):.1f} MB). "
                f"Maximum allowed size is {max_mb} MB."
            )

        suffix = Path(filename).suffix.lower()
        file_type = self._EXTENSION_MAP.get(suffix)
        if file_type is None:
            supported = ", ".join(self._EXTENSION_MAP.keys())
            raise ValueError(
                f"Unsupported file type '{suffix}'. Supported: {supported}"
            )

        parser_method = getattr(self, f"_parse_{file_type}")
        return parser_method(file_bytes, filename, file_type)

    # ------------------------------------------------------------------
    # Private per-format parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_csv(
        file_bytes: bytes, filename: str, file_type: str
    ) -> ParsedDocument:
        df = pd.read_csv(io.BytesIO(file_bytes))
        text_content = df.to_string(index=False)
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            text_content=text_content,
            dataframe=df,
        )

    @staticmethod
    def _parse_excel(
        file_bytes: bytes, filename: str, file_type: str
    ) -> ParsedDocument:
        # Read all sheets; concatenate into a single text block.
        xl = pd.ExcelFile(io.BytesIO(file_bytes))
        sheets: list[pd.DataFrame] = []
        text_parts: list[str] = []

        for sheet_name in xl.sheet_names:
            df_sheet = xl.parse(sheet_name)
            sheets.append(df_sheet)
            text_parts.append(f"=== Sheet: {sheet_name} ===\n{df_sheet.to_string(index=False)}")

        combined_df: Optional[pd.DataFrame]
        if not sheets:
            combined_df = pd.DataFrame()
        else:
            first_cols = list(sheets[0].columns)
            if all(list(df_s.columns) == first_cols for df_s in sheets[1:]):
                combined_df = pd.concat(sheets, ignore_index=True)
            else:
                # Sheets have different schemas — avoid silently merging them
                # into a misleading combined DataFrame.
                logger.warning(
                    "Excel file '%s' has sheets with different column schemas; "
                    "combined DataFrame is not available.",
                    filename,
                )
                combined_df = None
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            text_content="\n\n".join(text_parts),
            dataframe=combined_df,
        )

    @staticmethod
    def _parse_pdf(
        file_bytes: bytes, filename: str, file_type: str
    ) -> ParsedDocument:
        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            text_content="\n\n".join(text_parts),
            dataframe=None,
        )

    @staticmethod
    def _parse_word(
        file_bytes: bytes, filename: str, file_type: str
    ) -> ParsedDocument:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text_content = "\n".join(paragraphs)
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            text_content=text_content,
            dataframe=None,
        )
